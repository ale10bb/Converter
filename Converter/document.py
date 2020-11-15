# -*- coding: utf-8 -*-
import re
import os.path
import shutil
from docx import Document
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfFileMerger, PdfFileReader

from Converter.macros import _LOGGER, _RESOURCE_PATH

def read_XT17(src_file_path:str) -> dict:
    ''' 读取{src_file_path}文件，获取项目编号、项目名称、方案日期、撰写人。

    Args:
        src_file_path(str): 读取的文件

    Returns:
        dict: {'code': (str), 'name': (str),'date': (str),'leader': (str)}

    Raises:
        ValueError: 如果{src_file_path}不存在或格式非法
    '''
    # 校验src_file_path的格式是否有效
    if not os.path.isfile(src_file_path) or os.path.splitext(src_file_path)[1].lower() not in ['.doc', '.docx']:
        raise ValueError('Invalid word file.')
    # 自动转换doc，失败时由convert_latest()抛出异常
    if os.path.splitext(src_file_path)[1].lower() == '.doc':
        src_file_path = convert_latest(src_file_path)

    ret = {'code': '', 'name': '','date': '','leader': ''}
    document = Document(src_file_path)
    # 读项目编号
    ## 印象中所有项目编号都能在前几行读到
    for paragraph in document.paragraphs[0:5]:
        re_result = re.search('SHTEC20[0-9]{2}DSYS[0-9]{4}', paragraph.text)
        if re_result:
            ret['code'] = re_result[0]
            break
    ret['name'] = document.tables[0].cell(0, 1).text.strip()
    ret['date'] = document.tables[0].cell(3, 1).text.strip()
    ret['leader'] = document.tables[-5].cell(0, 1).text.strip()
    _LOGGER.debug('[Converter/document/read_XT17] ret = "{}"'.format(ret))

    return ret


def make_XT09(dst_dir_path:str, info:dict) -> str:
    ''' 根据传入的项目编号、项目名称、撰写人等，在{dst_dir_path}下生成XT09文档。

    Args:
        dst_dir_path(str): 输出文件的目录；
        info(dict): 需包含code、name、dict、leader、advices（均为str）；

    Returns:
        str: 生成的XT09文档的路径
    '''

    dst_file_path = os.path.join(dst_dir_path, 'RD-XT09测评方案审核意见单{}.docx'.format(info['code']))
    _LOGGER.info('[Converter/document/make_XT09] Generating "{}"'.format(dst_file_path))
    # ===== 写入XT09 =====
    document = Document(_RESOURCE_PATH(os.path.join('res', 'XT09', 'RD-XT09测评方案审核意见单.docx')))
    ## 项目编号(bold)
    document.paragraphs[0].add_run(info['code']).bold = True
    ## 项目名称
    document.tables[0].cell(1, 1).text = info['name']
    ## 方案制定人
    document.tables[0].cell(2, 1).text = info['leader']
    ## 审核意见
    document.tables[0].cell(4, 0).add_paragraph(info['advices'])
    ## 方案时间
    document.tables[0].cell(5, 0).add_paragraph(info['date']).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save(dst_file_path)
    return dst_file_path


def convert_PDF(src_file_path:str) -> str:
    ''' 调用Word，将{src_file_path}转换为PDF并存放在同目录。

    Args:
        src_file_path(str): 需转换的文件

    Returns:
        str: 转换完成的文件路径
    
    Raises:
        ValueError: 如果{src_file_path}不存在或格式非法
        OSError: 如果调用Word程序失败
    '''
    # 校验src_file_path的格式是否有效
    if os.path.isfile(src_file_path) and os.path.splitext(src_file_path)[1].lower() == '.pdf':
        return src_file_path
    if not os.path.isfile(src_file_path) or os.path.splitext(src_file_path)[1].lower() not in ['.doc', '.docx']:
        raise ValueError('Invalid word file.')
    dst_file_path = os.path.splitext(src_file_path)[0] + '.pdf'

    # 打开并关联Word，处理结束后放在后台不管，必要时会自动复用
    _LOGGER.debug('[Converter/document/convert_PDF] Dispatching Word.Application.')
    try:
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
    except:
        _LOGGER.error('[Converter/document/convert_PDF] Failed to dispatch Word.Application.')
        raise OSError('Failed to dispatch Word.Application.')
    # 转换doc文件
    _LOGGER.info('[Converter/document/convert_PDF] Converting "{}"'.format(src_file_path))
    document = word.Documents.Open(FileName=src_file_path, ReadOnly=True)
    document.SaveAs2(FileName=dst_file_path, FileFormat=17) # wdFormatPDF=17
    document.Close(SaveChanges=0)
    return dst_file_path


def merge_PDF(src_file_paths:list=[], dst_dir_path:str='') -> str:
    ''' 调用PyPDF2，将{src_file_paths}合并，输出到{dst_dir_path}下的output.pdf。

    Args:
        src_file_paths(list): 需合并的文件；
        dst_dir_path(str): 输出文件的目录，缺省时或目录无效时输出到第一个有效文件所在的目录；

    Returns:
        str: 合并完成的文件路径
    
    Raises:
        ValueError: 如果{src_file_paths}中没有任何有效文件
    '''
    # 校验src_file_paths中的文件是否有效
    valid_file_paths = []
    for src_file_path in src_file_paths:
        if os.path.isfile(src_file_path) and os.path.splitext(src_file_path)[1].lower() == '.pdf':
            valid_file_paths.append(src_file_path)
    _LOGGER.debug('[Converter/document/merge_PDF] valid_file_paths = "{}".'.format(valid_file_paths))
    # 校验dst_dir_path是否有效
    if not dst_dir_path or not os.path.isdir(dst_dir_path):
        dst_dir_path = os.path.dirname(valid_file_paths[0])
    dst_file_path = os.path.join(dst_dir_path, 'output.pdf')
    
    # 合并PDF
    if len(valid_file_paths) == 0:
        raise ValueError('Not enough PDFs.')
    elif len(valid_file_paths) == 1:
        shutil.copy(valid_file_paths[0], dst_file_path)
    else:
        merger = PdfFileMerger() 
        for src_file_path in valid_file_paths: 
            _LOGGER.debug('[Converter/document/merge_PDF] Reading "{}".'.format(src_file_path))
            with open(src_file_path, 'rb') as f: 
                merger.append(PdfFileReader(f), 'rb') 
        _LOGGER.info('[Converter/document/merge_PDF] Writing "{}".'.format(dst_file_path))
        with open(dst_file_path, 'wb') as f: 
            merger.write(f) 

    return dst_file_path


def convert_latest(src_file_path:str) -> str:
    ''' 调用Word，将{src_file_path}转换最新格式。

    Args:
        src_file_path(str): 需转换的文件

    Returns:
        str: 转换完成的文件路径
    
    Raises:
        ValueError: 如果{src_file_path}不存在或格式非法
        OSError: 如果调用Word程序失败
    '''
    # 校验src_file_path的格式是否有效
    if not os.path.isfile(src_file_path) or os.path.splitext(src_file_path)[1].lower() not in ['.doc', '.docx']:
        raise ValueError('Invalid word file.')
    dst_file_path = os.path.splitext(src_file_path)[0] + '.docx'

    # 打开并关联Word，处理结束后放在后台不管，必要时会自动复用
    _LOGGER.debug('[Converter/document/convert_latest] Dispatching Word.Application.')
    try:
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
    except:
        _LOGGER.error('[Converter/document/convert_latest] Failed to dispatch Word.Application.')
        raise OSError('Failed to dispatch Word.Application.')
    # 打开并关联文档
    document = word.Documents.Open(FileName=src_file_path)
    _LOGGER.info('[Converter/document/convert_latest] CompatibilityMode = {}.'.format(document.CompatibilityMode))
    # 将word转换为最新文件格式
    if document.CompatibilityMode < 15: # CompatibilityMode=11-14(旧版)
        _LOGGER.info('[Converter/document/convert_latest] Converting "{}".'.format(dst_file_path))
        document.Convert()
        document.Save()
    document.Close()
    return dst_file_path